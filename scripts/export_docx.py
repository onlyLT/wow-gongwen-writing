# -*- coding: utf-8 -*-
"""
export_docx.py — 通用公文 → Word(.docx) 导出器（GB/T 9704 近似版式）

数据驱动：吃一个 JSON 规格，产出排版好的 .docx。任意文种通用
（通知/请示/报告/批复/函/意见/纪要 + 专报/简报/总结/方案/讲话稿…）。

用法：
    python export_docx.py spec.json                 # 输出到 spec 里的 output
    python export_docx.py spec.json -o 目标.docx    # 覆盖输出路径
    python export_docx.py --demo demo.json          # 写一份示例 JSON 供参考

JSON 规格（所有字段可选，缺省即不出该要素）：
{
  "output": "outputs/xxx.docx",           // 输出路径（-o 可覆盖）
  "red_header": "武义县财政局文件",         // 红头·发文机关标志（红、居中、大字）
  "brief_header": "财政专报",              // 专报/简报报头（红、居中；与 red_header 二选一）
  "issue_no": "第 5 期",                   // 期号（简报/专报）
  "issuer_left": "武义县财政局编",          // 报头下左侧编发单位
  "issuer_right": "2026年7月2日",          // 报头下右侧日期
  "signer": "签发人：张三",                // 上行文签发人（右对齐）
  "wenhao": "武财〔2026〕5号",             // 发文字号（居中）
  "red_line": true,                        // 是否画红色分隔线（有红头/报头时默认 true）
  "title": "××关于××的报告",              // 标题（居中，二号）
  "zhusong": "金华市财政局：",             // 主送机关（顶格）
  "body": [                                // 正文块，按顺序渲染
    {"h1": "一、主要做法", "text": "……"}, // 一级标题（黑体）+ 仿宋同段正文（run-on）
    {"text": "普通正文段落，自动首行缩进2字"},
    {"h1": "二、存在的问题"},              // 独立标题行（无 text）
    {"h2": "（一）小标题", "text": "……"}, // 二级标题（楷体）+ 正文
    {"h3": "1．小节标题"},                 // 三级标题（仿宋加粗），如"1．基本支出预算"
    {"plain": "自由行", "align": "center", "font": "楷体", "size": 14, "no_indent": true}
  ],
  "signoff_name": "武义县财政局",           // 署名（右对齐）
  "signoff_date": "2026年7月2日",          // 成文日期（右对齐）
  "fuzhu": "（联系人：李四　电话：0579-…）", // 附注（首行缩进）
  "fonts": {"body": "仿宋", "h1": "黑体", "h2": "黑体",
            "title": "宋体", "header": "宋体"}  // 可覆盖字体名（缺省见下）
}

正文块类型：
  {"text": "..."}                  普通正文段（仿宋三号，首行缩进2字）
  {"h1": "一、…", "text": "..."}    一级标题（黑体三号）+ 可选同段正文
  {"h2": "（一）…", "text": "..."}  二级标题（楷体三号）+ 可选同段正文
  {"h3": "1．…", "text": "..."}     三级标题（仿宋三号加粗）+ 可选同段正文
  {"plain": "...", ...}            自由行，可带 align/font/size/bold/color/no_indent

字体缺省（GB/T 9704 标准层级，经实际发文核校）：标题=方正小标宋简体二号加粗，
一级标题=黑体三号，二级标题=楷体_GB2312 三号，三级标题=仿宋_GB2312 三号加粗，
正文/附注=仿宋_GB2312，红头=宋体加粗。可在 fonts 里覆盖任一层级
（需本机已安装；未安装时 Word 会自动替换，不影响生成）。

依赖：python-docx（pip install python-docx）
"""
import sys, json, argparse
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0x00, 0x00)
ALIGN = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
         "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}
# 层级用字（经实际公文核校）：标题=方正小标宋二号，一级=黑体三号，
# 二级=楷体三号，三级=仿宋三号加粗，正文=仿宋三号。未装字体时 Word 自动替换。
DEFAULT_FONTS = {"body": "仿宋_GB2312", "h1": "黑体", "h2": "楷体_GB2312",
                 "h3": "仿宋_GB2312", "title": "方正小标宋简体", "header": "宋体"}


def set_font(run, cn, size=16, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = cn
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn)
    if color is not None:
        run.font.color.rgb = color


def add_run(p, text, cn, size=16, bold=False, color=None):
    r = p.add_run(text)
    set_font(r, cn, size, bold, color)
    return r


def new_para(doc, align=None, line=29, space_after=0, first_indent=None):
    p = doc.add_paragraph()
    if align in ALIGN:
        p.alignment = ALIGN[align]
    pf = p.paragraph_format
    pf.line_spacing = Pt(line)
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    if first_indent:
        pf.first_line_indent = Pt(first_indent)
    return p


def red_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '24')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'FF0000')
    pbdr.append(bottom); pPr.append(pbdr)


def build(spec):
    fonts = dict(DEFAULT_FONTS); fonts.update(spec.get("fonts", {}))
    doc = Document()
    s = doc.sections[0]
    s.top_margin, s.bottom_margin = Cm(3.7), Cm(3.5)
    s.left_margin, s.right_margin = Cm(2.8), Cm(2.6)

    has_header = bool(spec.get("red_header") or spec.get("brief_header"))

    # —— 版头 ——
    if spec.get("red_header"):
        p = new_para(doc, "center", line=52, space_after=4)
        add_run(p, spec["red_header"], fonts["header"], 40, True, RED)
    if spec.get("brief_header"):
        p = new_para(doc, "center", line=46, space_after=2)
        add_run(p, spec["brief_header"], fonts["header"], 36, True, RED)
    if spec.get("issue_no"):
        p = new_para(doc, "center", line=26)
        add_run(p, spec["issue_no"], fonts["body"], 15)
    if spec.get("issuer_left") or spec.get("issuer_right"):
        p = new_para(doc, None, line=26)
        add_run(p, spec.get("issuer_left", ""), fonts["body"], 14)
        add_run(p, "　" * 16, fonts["body"], 14)
        add_run(p, spec.get("issuer_right", ""), fonts["body"], 14)
    if spec.get("signer"):
        p = new_para(doc, "right", line=28)
        add_run(p, spec["signer"], fonts["body"], 15)
    if spec.get("wenhao"):
        p = new_para(doc, "center", line=26, space_after=4)
        add_run(p, spec["wenhao"], fonts["body"], 15)
    if spec.get("red_line", has_header):
        red_line(doc)

    # —— 标题 ——
    if spec.get("title"):
        p = new_para(doc, "center", line=36, space_after=6)
        add_run(p, spec["title"], fonts["title"], 22, True)

    # —— 主送 ——
    if spec.get("zhusong"):
        p = new_para(doc, "left", line=29)
        add_run(p, spec["zhusong"], fonts["body"], 16)

    # —— 正文块 ——
    for blk in spec.get("body", []):
        if "plain" in blk:
            fi = None if blk.get("no_indent") else Pt(blk.get("size", 16) * 2)
            p = new_para(doc, blk.get("align"), line=blk.get("line", 29),
                         first_indent=(blk.get("size", 16) * 2 if not blk.get("no_indent") else None))
            color = RED if blk.get("color") == "red" else None
            add_run(p, blk["plain"], blk.get("font", fonts["body"]),
                    blk.get("size", 16), blk.get("bold", False), color)
            continue
        head_key = next((k for k in ("h1", "h2", "h3") if k in blk), None)
        if head_key:
            p = new_para(doc, "left", line=29, first_indent=32)
            add_run(p, blk[head_key], fonts[head_key], 16, bold=(head_key == "h3"))
            if blk.get("text"):
                add_run(p, blk["text"], fonts["body"], 16)
        elif "text" in blk:
            p = new_para(doc, "justify", line=29, first_indent=32)
            add_run(p, blk["text"], fonts["body"], 16)

    # —— 署名 / 日期 ——
    if spec.get("signoff_name"):
        p = new_para(doc, "right", line=29)
        add_run(p, spec["signoff_name"] + "　　　", fonts["body"], 16)
    if spec.get("signoff_date"):
        p = new_para(doc, "right", line=29)
        add_run(p, spec["signoff_date"] + "　　", fonts["body"], 16)

    # —— 附注 ——
    if spec.get("fuzhu"):
        p = new_para(doc, "left", line=26, first_indent=28)
        add_run(p, spec["fuzhu"], fonts["body"], 14)

    return doc


DEMO = {
    "output": "outputs/示例-通知.docx",
    "red_header": "××县人民政府办公室文件",
    "wenhao": "×政办发〔2026〕1号",
    "title": "××县人民政府办公室关于××工作的通知",
    "zhusong": "各乡镇人民政府，县直各单位：",
    "body": [
        {"text": "为切实做好××工作，根据《××》精神，现将有关事项通知如下。"},
        {"h1": "一、总体要求。", "text": "……"},
        {"h1": "二、重点任务。", "text": "……"},
        {"h1": "三、工作要求。", "text": "……"},
        {"text": "请遵照执行。"}
    ],
    "signoff_name": "××县人民政府办公室",
    "signoff_date": "2026年×月×日"
}


def main():
    ap = argparse.ArgumentParser(description="通用公文 → Word 导出器")
    ap.add_argument("spec", nargs="?", help="JSON 规格文件路径")
    ap.add_argument("-o", "--out", help="覆盖输出路径")
    ap.add_argument("--demo", metavar="PATH", help="写一份示例 JSON 到 PATH")
    args = ap.parse_args()

    if args.demo:
        with open(args.demo, "w", encoding="utf-8") as f:
            json.dump(DEMO, f, ensure_ascii=False, indent=2)
        print("DEMO written:", args.demo); return

    if not args.spec:
        ap.error("需要提供 JSON 规格文件（或用 --demo 生成示例）")

    with open(args.spec, "r", encoding="utf-8") as f:
        spec = json.load(f)
    out = args.out or spec.get("output")
    if not out:
        ap.error("未指定输出路径（JSON 的 output 或 -o）")

    import os
    d = os.path.dirname(os.path.abspath(out))
    if d and not os.path.isdir(d):
        os.makedirs(d, exist_ok=True)
    build(spec).save(out)
    print("SAVED:", out)


if __name__ == "__main__":
    main()
